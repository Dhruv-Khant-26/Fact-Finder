
import streamlit as st
import joblib
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import io
from fpdf import FPDF
import requests
import re
from urllib.parse import urlparse, parse_qs

# Load model and vectorizer
model = joblib.load('fake_news_model.pkl')
vectorizer = joblib.load('tfidf_vectorizer.pkl')

# Sample news articles (using dummy since True.csv is large)
sample_news = [
    """WASHINGTON (Reuters) - The head of a conservative Republican faction in the U.S. Congress, who voted this month for a huge expansion of the national debt to pay for tax cuts, called himself a "fiscal conservative" on Sunday and urged budget restraint in 2018. In keeping with a sharp pivot under way among Republicans, U.S. Representative Mark Meadows, speaking on CBS' "Face the Nation," drew a hard line on federal spending, which lawmakers are bracing to do battle over in January. When they return from the holidays on Wednesday, lawmakers will begin trying to pass a federal budget in a fight likely to be linked to other issues, such as immigration policy, even as the November congressional election campaigns approach in which Republicans will seek to keep control of Congress. President Donald Trump and his Republicans want a big budget increase in military spending, while Democrats also want proportional increases for non-defense "discretionary" spending on programs that support education, scientific research, infrastructure, public health and environmental protection. "The (Trump) administration has already been willing to say: 'We're going to increase non-defense discretionary spending ... by about 7 percent,'" Meadows, chairman of the small but influential House Freedom Caucus, said on the program. "Now, Democrats are saying that's not enough, we need to give the government a pay raise of 10 to 11 percent. For a fiscal conservative, I don't see where the rationale is. ... Eventually you run out of other people's money," he said. Meadows was among Republicans who voted in late December for their party's debt-financed tax overhaul, which is expected to balloon the federal budget deficit and add about $1.5 trillion over 10 years to the $20 trillion national debt. "It's interesting to hear Mark talk about fiscal responsibility," Democratic U.S. Representative Joseph Crowley said on CBS. Crowley said the Republican tax bill would require the  United States to borrow $1.5 trillion, to be paid off by future generations, to finance tax cuts for corporations and the rich. "This is one of the least ... fiscally responsible bills we've ever seen passed in the history of the House of Representatives. I think we're going to be paying for this for many, many years to come," Crowley said. Republicans insist the tax package, the biggest U.S. tax overhaul in more than 30 years,  will boost the economy and job growth. House Speaker Paul Ryan, who also supported the tax bill, recently went further than Meadows, making clear in a radio interview that welfare or "entitlement reform," as the party often calls it, would be a top Republican priority in 2018. In Republican parlance, "entitlement" programs mean food stamps, housing assistance, Medicare and Medicaid health insurance for the elderly, poor and disabled, as well as other programs created by Washington to assist the needy. Democrats seized on Ryan's early December remarks, saying they showed Republicans would try to pay for their tax overhaul by seeking spending cuts for social programs. But the goals of House Republicans may have to take a back seat to the Senate, where the votes of some Democrats will be needed to approve a budget and prevent a government shutdown. Democrats will use their leverage in the Senate, which Republicans narrowly control, to defend both discretionary non-defense programs and social spending, while tackling the issue of the "Dreamers," people brought illegally to the country as children. Trump in September put a March 2018 expiration date on the Deferred Action for Childhood Arrivals, or DACA, program, which protects the young immigrants from deportation and provides them with work permits. The president has said in recent Twitter messages he wants funding for his proposed Mexican border wall and other immigration law changes in exchange for agreeing to help the Dreamers. Representative Debbie Dingell told CBS she did not favor linking that issue to other policy objectives, such as wall funding. "We need to do DACA clean," she said.  On Wednesday, Trump aides will meet with congressional leaders to discuss those issues. That will be followed by a weekend of strategy sessions for Trump and Republican leaders on Jan. 6 and 7, the White House said. Trump was also scheduled to meet on Sunday with Florida Republican Governor Rick Scott, who wants more emergency aid. The House has passed an $81 billion aid package after hurricanes in Florida, Texas and Puerto Rico, and wildfires in California. The package far exceeded the $44 billion requested by the Trump administration. The Senate has not yet voted on the aid."""
]

# Function to create sample TXT file
def create_sample_txt():
    return '\n\n'.join(sample_news).encode('utf-8')

# Function to create sample PDF file
def create_sample_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for news in sample_news:
        pdf.multi_cell(0, 10, news)
        pdf.ln()
    return pdf.output(dest='S').encode('latin1')

# Function to create sample DOCX file
def create_sample_docx():
    doc = Document()
    for news in sample_news:
        doc.add_paragraph(news)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def download_from_google_drive(url):
    try:
        parsed = urlparse(url)
        if 'drive.google.com' not in parsed.netloc:
            return None, "Not a Google Drive link"

        file_id = None
        qs = parse_qs(parsed.query)
        if 'id' in qs:
            file_id = qs['id'][0]
        else:
            parts = parsed.path.split('/')
            if 'd' in parts:
                idx = parts.index('d')
                if len(parts) > idx + 1:
                    file_id = parts[idx + 1]

        if not file_id:
            return None, "Could not extract file id from link"

        session = requests.Session()
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        resp = session.get(download_url)

        # Handle large file confirmation token
        if resp.status_code == 200 and 'confirm=' in resp.text:
            m = re.search(r"confirm=([0-9A-Za-z_-]+)", resp.text)
            if m:
                confirm = m.group(1)
                resp = session.get(download_url + "&confirm=" + confirm)

        if resp.status_code != 200:
            return None, f"Failed to download file (status {resp.status_code})"

        return resp.content, None
    except Exception as e:
        return None, str(e)

# Set page config
st.set_page_config(page_title="Fake News Detection", layout="wide")

# Add background style
st.markdown(
    """
    <style>
    .stApp {
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Sidebar menu
with st.sidebar:
    st.header("Menu")
    st.subheader("Download Sample Files")
    st.download_button(
        label="Download TXT",
        data=create_sample_txt(),
        file_name="sample_news.txt",
        mime="text/plain"
    )
    st.download_button(
        label="Download PDF",
        data=create_sample_pdf(),
        file_name="sample_news.pdf",
        mime="application/pdf"
    )
    st.download_button(
        label="Download DOC",
        data=create_sample_docx(),
        file_name="sample_news.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Main content
st.title("Fake News Detection")

# Let user choose upload source
source = st.radio("Upload source", ("Local Upload", "Google Drive Link", "Paste Text"))

text = ""

if source == "Local Upload":
    uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'])

    if uploaded_file is not None:
        # Extract text based on file type
        try:
            if uploaded_file.type == "text/plain":
                text = uploaded_file.read().decode('utf-8')
            elif uploaded_file.type == "application/pdf":
                pdf_reader = PdfReader(uploaded_file)
                for page in pdf_reader.pages:
                    text += (page.extract_text() or "") + "\n"
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                text = '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            st.error(f"Error reading file: {e}")

        # Preview section
        st.subheader("File Preview")
        st.text_area("Extracted Text", text, height=300)

elif source == "Google Drive Link":
    gdrive_link = st.text_input("Paste Google Drive shareable link")
    st.markdown(
        '<a href="https://drive.google.com/drive/my-drive" target="_blank"><button style="padding:8px 12px;border-radius:6px">Open Google Drive</button></a>',
        unsafe_allow_html=True,
    )
    st.caption("Opens Google Drive in a new tab — download the file there and upload it here, or paste the shareable link.")
    if st.button("Fetch from Google Drive"):
        if gdrive_link.strip():
            content, err = download_from_google_drive(gdrive_link.strip())
            if err:
                st.error(f"Error fetching from Drive: {err}")
            else:
                # Try to parse as text, then PDF, then DOCX
                try:
                    text = content.decode('utf-8')
                except Exception:
                    parsed = False
                    try:
                        pdf_reader = PdfReader(io.BytesIO(content))
                        for page in pdf_reader.pages:
                            text += (page.extract_text() or "") + "\n"
                        parsed = True
                    except Exception:
                        try:
                            doc = Document(io.BytesIO(content))
                            text = '\n'.join([p.text for p in doc.paragraphs])
                            parsed = True
                        except Exception:
                            parsed = False

                    if not parsed and not text:
                        st.error("Could not parse file from Google Drive. Use a valid TXT, PDF, or DOCX file.")

                if text:
                    st.subheader("File Preview")
                    st.text_area("Extracted Text", text, height=300)

elif source == "Paste Text":
    text = st.text_area("Paste your news article here", height=300)

# Predict button (works for local upload, Google Drive, and pasted text)
if st.button("Predict"):
    if text.strip():
        text_vectorized = vectorizer.transform([text])
        prediction = model.predict(text_vectorized)[0]
        probability = model.predict_proba(text_vectorized)[0][1]

        if prediction == 1:
            st.error(f"Fake News (Probability: {probability:.2f})")
        else:
            st.success(f"Real News (Probability: {1-probability:.2f})")
    else:
        st.warning("No text extracted from the file.")